import re
import jieba
import time
import streamlit as st
from io import BytesIO, StringIO
from langdetect import detect
from yake import KeywordExtractor
from deep_translator import GoogleTranslator
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
import pdfplumber
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from textwrap import wrap
# 新增：Word文档生成库
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ===================== 全局页面配置 =====================
st.set_page_config(
    page_title="文献伴侣 | 华师大学生专属",
    layout="wide",
    initial_sidebar_state="collapsed"
)

hide_streamlit_style = """
<style>
#MainMenu, footer, header, .stToolbar, .deployButton {visibility: hidden !important;}
div, p, span, h1, h2, h3, h4, h5, h6, label, .stMarkdown {color: #1e293b !important;}
.stApp {background-color: #f8f9fa !important;}
.stCard {border-radius: 12px; padding: 20px; background: #ffffff; box-shadow: 0 2px 8px rgba(0,0,0,0.05); margin-bottom: 15px;}
.stButton > button {border-radius: 8px;}
textarea, .stTextInput > div > input {border-radius: 8px;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# 注册PDF字体
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
FONT_CN = 'STSong-Light'
FONT_SIZE = 11

# ===================== 会话初始化 =====================
def init_session():
    if "step" not in st.session_state:
        st.session_state.step = "menu"
    if "paper_text" not in st.session_state:
        st.session_state.paper_text = ""
    if "meta" not in st.session_state:
        st.session_state.meta = {"title":"", "author":"", "year":"", "journal":"", "volume":"", "pages":""}
    if "summary_core" not in st.session_state:
        st.session_state.summary_core = {"zh":"", "en":""}
    if "summary_detail" not in st.session_state:
        st.session_state.summary_detail = {"zh":"", "en":""}
    if "keywords" not in st.session_state:
        st.session_state.keywords = {"zh":[], "en":[]}
    if "lang" not in st.session_state:
        st.session_state.lang = ""
    if "paper_history" not in st.session_state:
        st.session_state.paper_history = []
    # 新增：摘要长度设置
    if "summary_lengths" not in st.session_state:
        st.session_state.summary_lengths = {"core": 100, "detail": 300}

init_session()

yake_ext = KeywordExtractor(n=3, top=3)
translator = GoogleTranslator

# ===================== 工具1：超长文本自动分段 =====================
def split_long_text(text, chunk_len=2000):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_len
        chunks.append(text[start:end])
        start = end
    return chunks

# ===================== 工具2：PDF文字提取 =====================
def extract_pdf_text(pdf_file):
    try:
        with pdfplumber.open(pdf_file) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"
        return text.strip()
    except Exception:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n\n"
        return text.strip()

# ===================== 工具3：增强元数据提取 =====================
def auto_extract_meta(text):
    meta = {"title":"", "author":"", "year":"", "journal":"", "volume":"", "pages":""}
    year_match = re.findall(r'19\d{2}|20\d{2}', text)
    if year_match:
        meta["year"] = year_match[0]
    page_match = re.findall(r'\d{1,4}[-–]\d{1,4}', text)
    if page_match:
        meta["pages"] = page_match[0]
    vol_match = re.search(r'Vol\.?\s*(\d+)|卷[:：]\s*(\d+)', text)
    if vol_match:
        meta["volume"] = vol_match.group(1) or vol_match.group(2)
    return meta

# ===================== 工具4：学术化摘要润色（去口语） =====================
def polish_zh_academic(text):
    text = re.sub(r'本文认为', '研究表明', text)
    text = re.sub(r'我们发现', '研究结果显示', text)
    text = re.sub(r'可以看出', '分析结果表明', text)
    text = re.sub(r'总而言之', '综上可知', text)
    text = re.sub(r'也就是说', '即', text)
    return text.strip()

def polish_en_academic(text):
    text = re.sub(r'we find', 'this study indicates', text, flags=re.I)
    text = re.sub(r'it can be seen', 'the analysis reveals', text, flags=re.I)
    text = re.sub(r'in a word', 'in conclusion', text, flags=re.I)
    return text.strip()

# ===================== 工具5：自定义长度学术摘要生成 =====================
def gen_academic_summary(text, lang, core_len=100, detail_len=300):
    chunks = split_long_text(text)
    core_list = []
    detail_list = []

    for chunk in chunks:
        if len(chunk) < 50:
            continue
        if lang == "en":
            parser = PlaintextParser.from_string(chunk, Tokenizer("english"))
            summarizer = LsaSummarizer()
            summary = summarizer(parser.document, 2)
            sentences = [str(s).strip() for s in summary if str(s).strip()]
            if sentences:
                core_list.append(sentences[0])
                detail_list.extend(sentences)
        else:
            sentences = re.split(r'[。！？；]', chunk)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 15]
            if sentences:
                core_list.append(sentences[0])
                detail_list.extend(sentences)

    # 合并并按长度截断
    core_raw = "。".join(core_list[:2]) if core_list else text[:core_len]
    detail_raw = "。".join(detail_list[:4]) if detail_list else core_raw
    
    # 按用户设定长度截取
    core_raw = core_raw[:core_len] + "..." if len(core_raw) > core_len else core_raw
    detail_raw = detail_raw[:detail_len] + "..." if len(detail_raw) > detail_len else detail_raw

    # 学术润色
    if lang == "zh":
        core_raw = polish_zh_academic(core_raw)
        detail_raw = polish_zh_academic(detail_raw)
    else:
        core_raw = polish_en_academic(core_raw)
        detail_raw = polish_en_academic(detail_raw)

    return core_raw, detail_raw

# ===================== 工具6：关键词提取 =====================
def gen_keywords(text):
    kw = yake_ext.extract_keywords(text)
    return [item[0] for item in kw[:3]]

# ===================== 工具7：中英互译 =====================
def translate_text(text, src, dest):
    if not text:
        return ""
    try:
        return translator(source=src, target=dest).translate(text)
    except:
        return "翻译失败，请检查网络"

# ===================== 工具8：三格式引用 =====================
def gen_all_citation(meta):
    author = meta["author"].strip() or "佚名"
    year = meta["year"].strip() or "不详"
    title = meta["title"].strip() or "无题文献"
    journal = meta["journal"].strip() or "未知期刊"
    volume = meta["volume"].strip()
    pages = meta["pages"].strip()

    apa = f"{author}. ({year}). {title}. *{journal}*, {volume}, {pages}."
    mla = f"{author}. \"{title}.\" *{journal}*, vol.{volume}, {year}, pp.{pages}."
    gb = f"{author}. {title}[J]. {journal}, {year}, {volume}:{pages}." if volume and pages else f"{author}. {title}[J]. {journal}, {year}."
    return apa, mla, gb

# ===================== 工具9：左右对照PDF =====================
def create_contrast_pdf(ori_text, tra_text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    w, h = A4
    left_w = w/2 - 30
    right_start = w/2 + 10
    line_h = 16
    start_y = h - 40

    c.setFont(FONT_CN, FONT_SIZE)
    c.drawString(20, start_y, "📄 原文")
    c.drawString(right_start, start_y, "📄 译文")
    start_y -= 25

    ori_paras = ori_text.split("\n")
    tra_paras = tra_text.split("\n")
    max_p = max(len(ori_paras), len(tra_paras))

    for idx in range(max_p):
        o_para = ori_paras[idx] if idx < len(ori_paras) else ""
        t_para = tra_paras[idx] if idx < len(tra_paras) else ""
        o_lines = wrap(o_para, width=45)
        t_lines = wrap(t_para, width=45)
        max_l = max(len(o_lines), len(t_lines))

        for i in range(max_l):
            if start_y < 40:
                c.showPage()
                c.setFont(FONT_CN, FONT_SIZE)
                start_y = h - 40
            o_line = o_lines[i] if i < len(o_lines) else ""
            t_line = t_lines[i] if i < len(t_lines) else ""
            c.drawString(20, start_y, o_line)
            c.drawString(right_start, start_y, t_line)
            start_y -= line_h
        start_y -= 8

    c.save()
    buffer.seek(0)
    return buffer

# ===================== 新增工具10：生成完整文献Word =====================
def create_literature_word(meta, summary_core, summary_detail, keywords, citations, ori_text, tra_text):
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('文献伴侣分析报告', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 元数据
    doc.add_heading('一、文献元数据', level=1)
    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.style = 'Table Grid'
    hdr_cells = meta_table.rows[0].cells
    hdr_cells[0].text = '标题'
    hdr_cells[1].text = meta["title"] or "未命名文献"
    hdr_cells = meta_table.rows[1].cells
    hdr_cells[0].text = '作者'
    hdr_cells[1].text = meta["author"] or "佚名"
    hdr_cells = meta_table.rows[2].cells
    hdr_cells[0].text = '发表信息'
    hdr_cells[1].text = f"{meta['year']} | {meta['journal']} | {meta['volume']}:{meta['pages']}"

    # 核心观点
    doc.add_heading('二、核心观点（双语）', level=1)
    doc.add_heading('英文', level=2)
    doc.add_paragraph(summary_core["en"])
    doc.add_heading('中文', level=2)
    doc.add_paragraph(summary_core["zh"])

    # 详细摘要
    doc.add_heading('三、详细摘要（双语）', level=1)
    doc.add_heading('英文', level=2)
    doc.add_paragraph(summary_detail["en"])
    doc.add_heading('中文', level=2)
    doc.add_paragraph(summary_detail["zh"])

    # 关键词
    doc.add_heading('四、关键论点（双语）', level=1)
    doc.add_paragraph(f"英文：{', '.join(keywords['en'])}")
    doc.add_paragraph(f"中文：{', '.join(keywords['zh'])}")

    # 引用格式
    doc.add_heading('五、参考文献格式', level=1)
    doc.add_paragraph(f"APA格式：{citations[0]}")
    doc.add_paragraph(f"MLA格式：{citations[1]}")
    doc.add_paragraph(f"GB/T 7714国标格式：{citations[2]}")

    # 保存到BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===================== 新增工具11：生成翻译全文Word =====================
def create_translation_word(ori_text, tra_text, lang):
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    src_lang = "中文" if lang == "zh" else "英文"
    dest_lang = "英文" if lang == "zh" else "中文"
    title = doc.add_heading(f'文献全文翻译（{src_lang}→{dest_lang}）', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 原文
    doc.add_heading('一、原文', level=1)
    for para in ori_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    # 译文
    doc.add_heading('二、译文', level=1)
    for para in tra_text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # 保存到BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ===================== 侧边栏：文献历史记录 =====================
with st.sidebar:
    st.markdown("### 📚 文献历史记录")
    if st.session_state.paper_history:
        for idx, item in enumerate(st.session_state.paper_history, 1):
            st.markdown(f"**{idx}. {item['title']}**")
            st.caption(f"{item['year']} | {item['journal']}")
        if st.button("清空历史记录"):
            st.session_state.paper_history = []
            st.rerun()
    else:
        st.caption("暂无分析记录")

# ===================== 主交互流程 =====================
st.markdown('<div class="stCard">', unsafe_allow_html=True)
st.title("📚 文献伴侣 | 华师大学生专属文献助手")
st.caption("超长文献分段解析 · 学术双语摘要 · 三格式引用 · 对照PDF · 历史记录 · 文档导出")
st.markdown('</div>', unsafe_allow_html=True)

# 主菜单
if st.session_state.step == "menu":
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("请选择使用方式")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("1️⃣ 粘贴论文文本", use_container_width=True):
            st.session_state.step = "input_text"
            st.rerun()
    with c2:
        if st.button("2️⃣ 上传文字型PDF", use_container_width=True):
            st.session_state.step = "input_pdf"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# 粘贴文本
elif st.session_state.step == "input_text":
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("粘贴论文全文（支持万字长文）")
    # 新增：摘要长度设置
    st.subheader("📏 摘要长度设置（字符数）")
    col1, col2 = st.columns(2)
    with col1:
        core_len = st.number_input("核心观点长度", min_value=50, max_value=500, value=st.session_state.summary_lengths["core"], step=10)
    with col2:
        detail_len = st.number_input("详细摘要长度", min_value=100, max_value=1000, value=st.session_state.summary_lengths["detail"], step=20)
    st.session_state.summary_lengths = {"core": core_len, "detail": detail_len}
    
    text = st.text_area("在此粘贴", height=320)
    c1, c2 = st.columns(2)
    with c1:
        if st.button("开始智能分析", use_container_width=True) and text.strip():
            st.session_state.paper_text = text.strip()
            st.session_state.step = "analyze"
            st.rerun()
    with c2:
        if st.button("返回主菜单", use_container_width=True):
            st.session_state.step = "menu"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# 上传PDF
elif st.session_state.step == "input_pdf":
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("上传文字版PDF（支持多页长篇文献）")
    # 新增：摘要长度设置
    st.subheader("📏 摘要长度设置（字符数）")
    col1, col2 = st.columns(2)
    with col1:
        core_len = st.number_input("核心观点长度", min_value=50, max_value=500, value=st.session_state.summary_lengths["core"], step=10)
    with col2:
        detail_len = st.number_input("详细摘要长度", min_value=100, max_value=1000, value=st.session_state.summary_lengths["detail"], step=20)
    st.session_state.summary_lengths = {"core": core_len, "detail": detail_len}
    
    pdf_file = st.file_uploader("选择PDF文件", type="pdf")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("解析并分析", use_container_width=True) and pdf_file:
            txt = extract_pdf_text(pdf_file)
            if txt:
                st.session_state.paper_text = txt
                st.session_state.step = "analyze"
                st.rerun()
            else:
                st.warning("解析失败！请使用可复制文字的PDF")
    with c2:
        if st.button("返回主菜单", use_container_width=True):
            st.session_state.step = "menu"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# 自动分析（长文本分段+学术摘要+自定义长度）
elif st.session_state.step == "analyze":
    with st.spinner("🔍 正在分段解析长文献、生成学术化双语摘要..."):
        text = st.session_state.paper_text
        core_len = st.session_state.summary_lengths["core"]
        detail_len = st.session_state.summary_lengths["detail"]
        
        raw_lang = detect(text)
        st.session_state.lang = "zh" if raw_lang in ["zh-cn", "zh-tw"] else "en"
        st.session_state.meta = auto_extract_meta(text)
        # 自定义长度学术摘要
        core_raw, detail_raw = gen_academic_summary(text, st.session_state.lang, core_len, detail_len)
        # 双语翻译
        if st.session_state.lang == "zh":
            st.session_state.summary_core["zh"] = core_raw
            st.session_state.summary_core["en"] = translate_text(core_raw, "zh", "en")
            st.session_state.summary_detail["zh"] = detail_raw
            st.session_state.summary_detail["en"] = translate_text(detail_raw, "zh", "en")
        else:
            st.session_state.summary_core["en"] = core_raw
            st.session_state.summary_core["zh"] = translate_text(core_raw, "en", "zh")
            st.session_state.summary_detail["en"] = detail_raw
            st.session_state.summary_detail["zh"] = translate_text(detail_raw, "en", "zh")
        # 关键词
        kw_raw = gen_keywords(text)
        if st.session_state.lang == "zh":
            st.session_state.keywords["zh"] = kw_raw
            st.session_state.keywords["en"] = [translate_text(k, "zh", "en") for k in kw_raw]
        else:
            st.session_state.keywords["en"] = kw_raw
            st.session_state.keywords["zh"] = [translate_text(k, "en", "zh") for k in kw_raw]
    st.session_state.step = "result"
    st.rerun()

# 结果展示 + 存入历史记录 + 文档导出
elif st.session_state.step == "result":
    text = st.session_state.paper_text
    meta = st.session_state.meta
    core_len = st.session_state.summary_lengths["core"]
    detail_len = st.session_state.summary_lengths["detail"]

    # 存入历史记录
    record = {
        "title": meta["title"] or "未命名文献",
        "year": meta["year"] or "未知年份",
        "journal": meta["journal"] or "未知期刊",
        "time": time.strftime("%m-%d %H:%M")
    }
    if record not in st.session_state.paper_history:
        st.session_state.paper_history.append(record)

    # 元数据编辑
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("🔖 文献元数据（可手动补充）")
    c1,c2,c3 = st.columns(3)
    with c1:
        meta["title"] = st.text_input("论文标题", value=meta["title"])
        meta["author"] = st.text_input("作者（多作者用分号分隔）", value=meta["author"])
    with c2:
        meta["year"] = st.text_input("发表年份", value=meta["year"])
        meta["journal"] = st.text_input("期刊/会议名称", value=meta["journal"])
    with c3:
        meta["volume"] = st.text_input("卷号", value=meta["volume"])
        meta["pages"] = st.text_input("页码", value=meta["pages"])
    st.session_state.meta = meta
    st.markdown('</div>', unsafe_allow_html=True)

    # 学术双语摘要
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("💡 学术核心观点（双语）")
    st.caption(f"（自定义长度：{core_len}字符）")
    st.markdown(f"**英文**：{st.session_state.summary_core['en']}")
    st.markdown(f"**中文**：{st.session_state.summary_core['zh']}")
    st.divider()
    st.subheader("📝 学术详细摘要（双语）")
    st.caption(f"（自定义长度：{detail_len}字符）")
    st.markdown(f"**英文**：{st.session_state.summary_detail['en']}")
    st.markdown(f"**中文**：{st.session_state.summary_detail['zh']}")
    st.markdown('</div>', unsafe_allow_html=True)

    # 关键词
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("🔑 三大关键论点")
    st.markdown(f"**英文**：{', '.join(st.session_state.keywords['en'])}")
    st.markdown(f"**中文**：{', '.join(st.session_state.keywords['zh'])}")
    st.markdown('</div>', unsafe_allow_html=True)

    # 三种引用
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("📎 三格式参考文献（华师毕业论文直接用）")
    apa, mla, gb = gen_all_citation(meta)
    st.markdown(f"**APA格式**：{apa}")
    st.markdown(f"**MLA格式**：{mla}")
    st.markdown(f"**GB/T 7714国标格式**：{gb}")
    st.markdown('</div>', unsafe_allow_html=True)

    # 对照PDF下载 + 新增文档导出功能
    st.markdown('<div class="stCard">', unsafe_allow_html=True)
    st.subheader("📄 文档导出中心（PDF+Word）")
    
    # 全文翻译
    if st.session_state.lang == "zh":
        full_tra = translate_text(text, "zh", "en")
    else:
        full_tra = translate_text(text, "en", "zh")
    
    # 1. 左右对照PDF下载
    pdf_buf = create_contrast_pdf(text, full_tra)
    st.download_button(
        label="📥 下载 双语对照PDF",
        data=pdf_buf,
        file_name="文献双语对照_华师专用.pdf",
        mime="application/pdf",
        use_container_width=True
    )
    
    # 2. 新增：完整文献分析报告Word下载
    citations = (apa, mla, gb)
    word_buf = create_literature_word(meta, st.session_state.summary_core, 
                                     st.session_state.summary_detail, st.session_state.keywords,
                                     citations, text, full_tra)
    st.download_button(
        label="📥 下载 完整文献分析报告（Word）",
        data=word_buf,
        file_name="文献分析报告_华师专用.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    
    # 3. 新增：整篇文献翻译Word下载
    trans_word_buf = create_translation_word(text, full_tra, st.session_state.lang)
    src_lang = "中文" if st.session_state.lang == "zh" else "英文"
    dest_lang = "英文" if st.session_state.lang == "zh" else "中文"
    st.download_button(
        label=f"📥 下载 整篇文献翻译（{src_lang}→{dest_lang} Word）",
        data=trans_word_buf,
        file_name=f"文献全文翻译_{src_lang}转{dest_lang}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.markdown('</div>', unsafe_allow_html=True)

    # 重置
    if st.button("🔄 重新分析新文献", use_container_width=True):
        for k in list(st.session_state.keys()):
            if k != "paper_history":
                del st.session_state[k]
        st.rerun()
